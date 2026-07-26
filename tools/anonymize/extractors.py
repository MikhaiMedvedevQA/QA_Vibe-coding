"""Извлечение текста и изображений из входных форматов.

Возвращает нормализованное представление:
- текстовые форматы (txt/pdf/docx/doc) -> DocumentContent (упорядоченный список блоков);
- Excel -> ExcelContent (рабочая книга openpyxl + список медиа из xl/media).

Изображения НЕ анонимизируются: выносятся как есть, привязка к тексту
формируется в writers.py через индекс абзаца-плейсхолдера.
"""

import re
import subprocess
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
from openpyxl import load_workbook

from config import SUPPORTED_TEXT_EXTS, SUPPORTED_EXCEL_EXTS, SUPPORTED_EXTS


@dataclass
class TextBlock:
    text: str


@dataclass
class ImageBlock:
    asset_id: str       # img-001, img-002 ...
    data: bytes
    ext: str            # png, jpg, jpeg ...
    src_page: int | None = None
    note: str = ""


@dataclass
class EmbeddedBlock:
    """Внедрённый файл (OLE-объект: Excel/Word-шаблон и т.п.). Не анонимизируется,
    выносится как есть; в текст вставляется ссылка с именем файла."""
    asset_id: str       # att-001, att-002 ...
    data: bytes
    original_filename: str  # имя из word/embeddings/, напр. Microsoft_Excel_97-2003_Worksheet.xls
    ext: str                # xls, docx, bin ...
    prog_id: str = ""       # ProgID, напр. Excel.Sheet.8
    note: str = ""


@dataclass
class TableBlock:
    """Таблица документа как список строк ячеек. writers рендерит её в
    markdown-таблицу (первая строка — заголовок)."""
    rows: list  # list[list[str]]


@dataclass
class DocumentContent:
    source: Path
    source_format: str  # txt | pdf | docx | doc
    blocks: list = field(default_factory=list)


@dataclass
class ExcelMedia:
    asset_id: str
    data: bytes
    ext: str
    src_path: str  # путь внутри xlsx, напр. xl/media/image1.png


def _norm(s: str) -> str:
    """Нормализация извлечённого текста: NBSP -> пробел, soft hyphen -> дефис,
    zero-width пробелы удаляются. PDF-экстракторы часто отдают \xa0/\xad вместо
    обычных пробелов/дефисов, из-за чего regex и справочники промахиваются."""
    if not s:
        return s
    return s.replace("\xa0", " ").replace(" ", " ").replace(" ", " ") \
            .replace("\xad", "-").replace("‐", "-").replace("‑", "-") \
            .replace("​", "").replace("‌", "").replace("‍", "")


@dataclass
class ExcelContent:
    source: Path
    workbook: object  # openpyxl Workbook
    media: list = field(default_factory=list)


# --------------------------------------------------------------------------
# TXT
# --------------------------------------------------------------------------

def extract_txt(path: Path) -> DocumentContent:
    raw = path.read_text(encoding="utf-8", errors="replace")
    blocks = [TextBlock(text=_norm(line)) for line in raw.splitlines() if line.strip()]
    return DocumentContent(source=path, source_format="txt", blocks=blocks)


# --------------------------------------------------------------------------
# PDF (PyMuPDF)
# --------------------------------------------------------------------------

def extract_pdf(path: Path) -> DocumentContent:
    doc = fitz.open(path)
    blocks: list = []
    img_counter = 0
    for page_idx in range(len(doc)):
        page = doc[page_idx]
        page_no = page_idx + 1
        text = page.get_text("text") or ""
        # Разбиваем страницу на абзацы по непустым строкам.
        for line in text.splitlines():
            if line.strip():
                blocks.append(TextBlock(text=_norm(line.strip())))
        # Картинки страницы.
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base = doc.extract_image(xref)
            except Exception:
                continue
            img_counter += 1
            blocks.append(ImageBlock(
                asset_id=f"img-{img_counter:03d}",
                data=base["image"],
                ext=base.get("ext", "png"),
                src_page=page_no,
            ))
    doc.close()
    return DocumentContent(source=path, source_format="pdf", blocks=blocks)


# --------------------------------------------------------------------------
# DOCX (python-docx)
# --------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _para_text(paragraph) -> str:
    """Весь видимый текст абзаца, включая текст внутри <w:ins> (вставленные правки
    при включённом рецензировании). python-docx paragraph.text такой текст пропускает.
    Удалённый текст (w:delText) не берём — его в финальном документе нет."""
    texts = paragraph._p.findall(f".//{{{_W_NS}}}t")
    return "".join((t.text or "") for t in texts)


def _docx_image_rids(paragraph) -> list[str]:
    """Найти rId встроенных изображений в абзаце (через a:blip @r:embed)."""
    p_xml = paragraph._p
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    rids = []
    for blip in p_xml.findall(".//a:blip", ns):
        rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rid:
            rids.append(rid)
    return rids


# o:OLEObject — внедрённый файл (Excel/Word-шаблон и т.п.).
_OO_NS = "urn:schemas-microsoft-com:office:office"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _docx_embedded_objects(paragraph) -> list[tuple[str, str]]:
    """Вернуть [(rid, prog_id), ...] для OLE-объектов в абзаце."""
    out = []
    for ole in paragraph._p.findall(f".//{{{_OO_NS}}}OLEObject"):
        rid = ole.get(f"{{{_R_NS}}}id")
        prog = ole.get("ProgID") or ""
        if rid:
            out.append((rid, prog))
    return out


def extract_docx(path: Path, document_part=None) -> DocumentContent:
    document = docx.Document(str(path)) if document_part is None else document_part
    blocks: list = []
    counters = {"img": 0, "emb": 0}

    def emit_paragraph(paragraph):
        txt = _para_text(paragraph).strip()
        if txt:
            blocks.append(TextBlock(text=_norm(txt)))
        for rid in _docx_image_rids(paragraph):
            try:
                part = document.part.related_parts[rid]
                img_bytes = part.blob
                ext = part.partname.ext.lstrip(".") if hasattr(part.partname, "ext") else "png"
            except Exception:
                continue
            counters["img"] += 1
            blocks.append(ImageBlock(
                asset_id=f"img-{counters['img']:03d}",
                data=img_bytes,
                ext=ext or "png",
                src_page=None,
            ))
        for rid, prog_id in _docx_embedded_objects(paragraph):
            try:
                part = document.part.related_parts[rid]
                data = part.blob
                partname = str(part.partname)  # /word/embeddings/<file>
                filename = partname.rsplit("/", 1)[-1]
                ext = filename.rsplit(".", 1)[-1] if "." in filename else "bin"
            except Exception:
                continue
            counters["emb"] += 1
            blocks.append(EmbeddedBlock(
                asset_id=f"att-{counters['emb']:03d}",
                data=data,
                original_filename=filename,
                ext=ext,
                prog_id=prog_id,
            ))

    def emit_table(table):
        rows = []
        for row in table.rows:
            cells = [_norm(" ".join(_para_text(p) for p in c.paragraphs).strip())
                     for c in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            blocks.append(TableBlock(rows=rows))

    # Идём по телу документа в XML-порядке, чередуя параграфы и таблицы.
    # document.paragraphs / document.tables отдельно теряют порядок — таблицы
    # скапливались бы в конце, ломая привязку вложений/картинок к месту в тексте.
    body = document.element.body
    p_tag = f"{{{_W_NS}}}p"
    tbl_tag = f"{{{_W_NS}}}tbl"
    for child in body.iterchildren():
        tag = child.tag
        if tag == p_tag:
            emit_paragraph(Paragraph(child, document))
        elif tag == tbl_tag:
            emit_table(Table(child, document))
    return DocumentContent(source=path, source_format="docx", blocks=blocks)


# --------------------------------------------------------------------------
# DOC (legacy) — конвертация pandoc -> docx во временный файл
# --------------------------------------------------------------------------

def extract_doc(path: Path) -> DocumentContent:
    """Legacy .doc: текст через antiword. Картинки не извлекаются (ограничение antiword);
    для полноценной обработки сохраните файл как .docx."""
    # antiword с mapping UTF-8 для корректной кириллицы; fallback на дефолт.
    for cmd in (["antiword", "-m", "UTF-8.txt", str(path)], ["antiword", str(path)]):
        try:
            res = subprocess.run(cmd, check=True, capture_output=True)
            text = res.stdout.decode("utf-8", errors="replace")
            break
        except FileNotFoundError:
            raise RuntimeError(
                "antiword не найден в PATH. Для обработки legacy .doc установите antiword, "
                "либо сохраните файл как .docx."
            )
        except subprocess.CalledProcessError:
            continue
    else:
        raise RuntimeError(f"antiword не смог прочитать {path.name}. Сохраните файл как .docx.")
    blocks = [TextBlock(text=_norm(line)) for line in text.splitlines() if line.strip()]
    return DocumentContent(source=path, source_format="doc", blocks=blocks)


# --------------------------------------------------------------------------
# XLSX (openpyxl + zipfile для медиа)
# --------------------------------------------------------------------------

def extract_xlsx(path: Path) -> ExcelContent:
    wb = load_workbook(path, data_only=False)
    media: list[ExcelMedia] = []
    counter = 0
    # Изображения лежат в xl/media/ внутри xlsx (это zip-архив).
    try:
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if name.startswith("xl/media/") and not name.endswith("/"):
                    counter += 1
                    ext = name.rsplit(".", 1)[-1].lower() if "." in name else "png"
                    media.append(ExcelMedia(
                        asset_id=f"img-{counter:03d}",
                        data=zf.read(name),
                        ext=ext,
                        src_path=name,
                    ))
    except zipfile.BadZipFile:
        pass
    return ExcelContent(source=path, workbook=wb, media=media)


# --------------------------------------------------------------------------
# Диспетчер
# --------------------------------------------------------------------------

def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTS


def extract(path: Path):
    """Вернуть DocumentContent (текстовые форматы) или ExcelContent (xlsx)."""
    ext = path.suffix.lower()
    if ext in SUPPORTED_EXCEL_EXTS:
        return extract_xlsx(path)
    if ext == ".txt":
        return extract_txt(path)
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".doc":
        return extract_doc(path)
    raise ValueError(f"Неподдерживаемое расширение: {ext}")